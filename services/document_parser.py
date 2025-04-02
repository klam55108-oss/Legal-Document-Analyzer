"""
Document parsing service for extracting text from various document formats.
"""
import os
import logging
import re
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional, Union

import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Define allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'rtf'}

def is_allowed_file(filename):
    """
    Check if a file has an allowed extension.
    
    Args:
        filename: The name of the file to check
        
    Returns:
        True if the file extension is allowed, False otherwise
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class DocumentParser:
    """Service for parsing different document formats."""
    
    def __init__(self):
        """Initialize the document parser."""
        logger.info("Document parser initialized")
    
    def convert_to_txt(self, file_path: str) -> str:
        """
        Convert a document to text file format.
        
        Args:
            file_path: Path to the original document file
            
        Returns:
            Path to the new txt file, or empty string if conversion failed
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found for conversion: {file_path}")
            return ""
        
        # Extract the text content from the document
        text_content = self.parse_document(file_path)
        
        if not text_content or len(text_content) < 10:
            logger.error(f"Failed to extract meaningful text from {file_path}")
            return ""
        
        # Create a new txt file path
        original_dir = os.path.dirname(file_path)
        original_name = os.path.basename(file_path)
        base_name = original_name.rsplit('.', 1)[0]
        new_path = os.path.join(original_dir, f"{base_name}.txt")
        
        # Write the text content to the file
        try:
            with open(new_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(text_content)
            logger.info(f"Successfully converted {file_path} to {new_path}")
            return new_path
        except Exception as e:
            logger.error(f"Error writing text file: {str(e)}")
            return ""
        
    def parse_document(self, file_path: str) -> str:
        """
        Parse a document and extract its text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            The extracted text content
        """
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return "Document file not found. It may have been moved or deleted."
            
        # Check file size
        try:
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                logger.warning(f"Empty file: {file_path}")
                return "The uploaded document is empty (0 bytes)."
                
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                logger.warning(f"File too large: {file_path} ({file_size / (1024*1024):.2f} MB)")
                return "The document is too large to process efficiently. Please upload a smaller document (< 50MB)."
        except Exception as size_err:
            logger.error(f"Error checking file size: {str(size_err)}")
            
        # Get file extension
        _, ext = os.path.splitext(file_path.lower())
        
        try:
            # Route to appropriate parser based on extension
            if ext == '.pdf':
                return self._parse_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif ext == '.txt':
                return self._parse_txt(file_path)
            elif ext == '.rtf':
                # For RTF files, convert to text and treat as TXT
                logger.info(f"RTF file detected, using basic text extraction: {file_path}")
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                        text = file.read()
                    return self._clean_text(text)
                except Exception as rtf_err:
                    logger.error(f"Error parsing RTF: {str(rtf_err)}")
                    return "RTF file format could not be parsed correctly. Please convert to PDF or DOCX and try again."
            else:
                logger.error(f"Unsupported file format: {ext}")
                return f"Unsupported file format: {ext}. Please upload a PDF, DOCX, DOC, or TXT file."
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            return f"Error analyzing document: {str(e)}. Please try a different document or format."
            
    def _parse_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            The extracted text content
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if the PDF has pages
                if len(pdf_reader.pages) == 0:
                    logger.warning(f"PDF has no pages: {file_path}")
                    return "This document appears to be empty or could not be properly parsed."
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        
                        # If page is empty, add a placeholder
                        if not page_text or page_text.strip() == "":
                            page_text = f"[Page {page_num + 1} appears to be empty or contains only images]"
                            
                        text += page_text + "\n\n"
                    except Exception as page_error:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(page_error)}")
                        text += f"[Error extracting text from page {page_num + 1}]\n\n"
                        
            # If no text was extracted, provide a default message
            if not text or text.strip() == "":
                logger.warning(f"No text could be extracted from PDF: {file_path}")
                return "This document appears to contain only images or could not be parsed properly. Please convert it to a searchable PDF and try again."
                    
            # Clean up text
            text = self._clean_text(text)
            
            # If after cleaning there's still no useful text, use sample text
            if not text or len(text) < 20:  # Arbitrary minimum threshold
                logger.warning(f"After cleaning, PDF has insufficient text: {file_path}")
                text = "This legal document appears to be in a format that couldn't be fully parsed. It may contain scanned images of text rather than selectable text. Please consider using a document with searchable text."
            
            return text
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            # Return a fallback message instead of raising an exception
            return "There was an error parsing this document. It may be password-protected, corrupted, or in an unsupported format. Please try another document."
            
    def _parse_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            The extracted text content
        """
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            if not paragraphs:
                logger.warning(f"No paragraph text found in DOCX: {file_path}")
                paragraphs = ["[Document contains no paragraph text]"]
                
            text = "\n".join(paragraphs)
            
            # Extract text from tables
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text += cell.text + "\n"
                    table_text += "\n"
                table_text += "\n"
            
            # Add table text if any was found
            if table_text.strip():
                text += "\n\n" + table_text
                
            # Clean up text
            text = self._clean_text(text)
            
            # If no meaningful text was extracted
            if not text or len(text) < 20:
                logger.warning(f"Insufficient text extracted from DOCX: {file_path}")
                text = "This document appears to contain minimal text or could not be parsed properly. It may be corrupt or contain only images or formatting."
            
            return text
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return "There was an error parsing this document. It may be corrupted or in an unsupported format. Please try another document."
            
    def _parse_txt(self, file_path: str) -> str:
        """
        Extract text from a TXT file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            The extracted text content
        """
        try:
            # First try UTF-8
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            except UnicodeDecodeError:
                # If UTF-8 fails, try with error replacement
                with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                    text = file.read()
                logger.warning(f"Used fallback encoding for TXT file: {file_path}")
                
            # Clean up text
            text = self._clean_text(text)
            
            # Check if we have any meaningful content
            if not text or len(text.strip()) < 10:
                logger.warning(f"Empty or very small TXT file: {file_path}")
                text = "This text file appears to be empty or contains very little content."
            
            return text
        except Exception as e:
            logger.error(f"Error parsing TXT: {str(e)}")
            return "There was an error reading this text file. It may be corrupted or inaccessible. Please try another document."
            
    def _clean_text(self, text: str) -> str:
        """
        Clean up extracted text.
        
        Args:
            text: The text to clean
            
        Returns:
            The cleaned text
        """
        if not text:
            return ""
            
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix line breaks
        text = re.sub(r'(\w) *\n *(\w)', r'\1 \2', text)
        
        # Restore paragraph breaks
        text = re.sub(r'(\.) *\n *([A-Z])', r'.\n\n\2', text)
        
        return text.strip()
        
# Create a singleton instance
document_parser = DocumentParser()